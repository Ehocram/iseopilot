#!/usr/bin/env python3
"""
folder_index.py — Indice incrementale + ricerca per la modalità "Base dati"
Sviluppato da Marco Bonometti

PERCHÉ ESISTE
-------------
La vecchia `read_kb_directory()` leggeva l'INTERA cartella (anche di rete) ad
ogni refresh, ne concatenava il testo in un'unica stringa e la iniettava nel
prompt. Conseguenze: (1) lentissima su share di rete (riapre e ri-parsa ogni
PDF/Word/Excel), (2) imprecisa, perché il modello riceveva solo i primi
~6.000/18.000 caratteri della concatenazione — cioè i primi file in ordine
ALFABETICO, non quelli pertinenti alla domanda.

Questo modulo trasforma la "Base dati" in una vera RICERCA:
  • Indice FTS5 (SQLite, libreria standard — nessuna dipendenza nuova).
  • Aggiornamento INCREMENTALE: ri-parsa solo i file nuovi/modificati
    (confronto mtime+size); i file invariati non vengono più riaperti.
  • Ricerca BM25 a tempo di query: al modello arrivano solo i passaggi
    pertinenti → più veloce e più preciso.
  • Walk con os.scandir (meno round-trip su SMB) e robusto ai percorsi di rete
    irraggiungibili (try/except per file, niente crash).
  • Re-rank semantico OPZIONALE (sentence-transformers) se già installato.

SICUREZZA (CISO)
----------------
L'indice è un file SQLite locale in ~/Documents/ChatAssistant/folder_index/ e
contiene il TESTO dei documenti: eredita la riservatezza della sorgente. Per
share di rete con contenuti riservati o condivisi tra più utenti resta
preferibile l'istanza ChromaDB CENTRALIZZATA (indicizzazione e accessi governati
una volta sola). L'indicizzatore legge i file con le credenziali dell'utente
stesso: nessun account di servizio, nessuna elevazione di privilegi.
"""
import os
import re
import time
import hashlib
import sqlite3
import datetime as _dt
from pathlib import Path


def _dbg(msg):
    """Log leggero sullo stesso file del resto dell'app (~/chat_assistant_debug.txt)."""
    try:
        with open(Path.home() / "chat_assistant_debug.txt", "a", encoding="utf-8") as f:
            f.write(f"[{_dt.datetime.now():%H:%M:%S}] [folder_index] {msg}\n")
    except Exception:
        pass


# ── Dove vive l'indice (stesso albero degli altri dati utente) ──
INDEX_DIR = Path.home() / "Documents" / "ChatAssistant" / "folder_index"
INDEX_DIR.mkdir(parents=True, exist_ok=True)

# Estensioni indicizzate (le immagini si saltano: niente da cercare nel testo)
SUPPORTED_EXT = {".txt", ".md", ".py", ".json", ".xml", ".csv",
                 ".pdf", ".docx", ".xlsx", ".xls"}

MAX_FILE_BYTES = 25 * 1024 * 1024     # salta file enormi (> 25 MB)
MAX_TEXT_CHARS = 400_000              # tetto testo estratto per file
CHUNK_WORDS    = 220                  # granularità passaggio (più piccola = più precisa)
CHUNK_OVERLAP  = 40

# Versione della LOGICA DI ESTRAZIONE del testo. Quando cambia (es. ora i .docx
# includono anche il testo delle TABELLE), l'indice esistente contiene testo
# estratto con la vecchia logica e i file NON risultano modificati (mtime/size
# invariati) → l'update incrementale non li ri-parserebbe. Cambiando questa
# costante, update() rigenera l'indice da zero una volta sola.
EXTRACTOR_VERSION = "2"  # v2: estrazione .docx con tabelle + intestazioni/piè di pagina

# Stopwords IT/EN per costruire la query FTS (non per filtrare i documenti)
_STOP = {
    "il","lo","la","gli","le","un","uno","una","di","da","in","con","su","per",
    "tra","fra","e","o","ma","se","non","che","chi","come","quando","dove","cosa",
    "qual","quali","quanto","quanti","quante","mi","ti","ci","vi","si","ne","ho",
    "ha","hai","hanno","sono","sei","siamo","siete","era","del","della","dei",
    "degli","delle","al","alla","ai","agli","alle","nel","nella","nei","negli",
    "nelle","sul","sulla","sui","sugli","sulle","dal","dalla","dai","mie","miei",
    "the","an","is","are","was","were","been","have","has","had","of","to","for",
    "and","or","in","on","with","dammi","dimmi","parlami","mostrami","quale",
    "questo","questa","questi","queste",
    # interrogative e funzionali inglesi (per utenti EN). NOTA: "it" è escluso
    # di proposito: in ISEO "IT" è il nome del reparto e deve restare cercabile.
    "what","which","who","whom","where","when","how","why","does","did","will",
    "would","could","should","can","may","might","must","show","give","tell",
    "list","find","please","about","from","that","this","these","those","there",
    "here","your","our","their","them","all","any","some","much","many",
}


# ── FTS5 disponibile? (in CPython standard sì; lo verifichiamo comunque) ──
def fts5_available() -> bool:
    try:
        c = sqlite3.connect(":memory:")
        c.execute("CREATE VIRTUAL TABLE _t USING fts5(x)")
        c.close()
        return True
    except Exception:
        return False


def _db_path_for(kb_dir: str) -> Path:
    h = hashlib.sha1(str(Path(kb_dir).resolve()).encode("utf-8")).hexdigest()[:16]
    return INDEX_DIR / f"kb_{h}.db"


def chunk_text(text: str, size: int = CHUNK_WORDS, overlap: int = CHUNK_OVERLAP):
    if not text or not text.strip():
        return []
    words = text.split()
    out, i = [], 0
    while i < len(words):
        c = " ".join(words[i:i + size])
        if c.strip():
            out.append(c)
        i += size - overlap
        if i >= len(words):
            break
    return out


def _docx_full_text(path: Path) -> str:
    """Estrae TUTTO il testo da un .docx: paragrafi + celle di TABELLA (anche
    tabelle annidate) + intestazioni/piè di pagina. python-docx con .paragraphs
    NON legge il testo dentro le tabelle: i documenti ISEO sono basati su template
    a tabelle, quindi senza questo il corpo del documento andrebbe perso e
    resterebbe indicizzato solo il titolo."""
    from docx import Document
    doc = Document(str(path))
    out = []

    def _walk(container):
        # container = Document, cella, header o footer: tutti hanno .paragraphs e .tables
        for p in getattr(container, "paragraphs", []):
            t = p.text
            if t and t.strip():
                out.append(t)
        for tbl in getattr(container, "tables", []):
            for row in tbl.rows:
                for cell in row.cells:
                    _walk(cell)  # ricorsivo → gestisce tabelle dentro celle

    _walk(doc)
    # intestazioni e piè di pagina (saltando quelli "ereditati" per non duplicare)
    for sec in doc.sections:
        for hf in (sec.header, sec.footer):
            try:
                if getattr(hf, "is_linked_to_previous", False):
                    continue
                _walk(hf)
            except Exception:
                pass
    return "\n".join(out)


def _extract_text(path: Path, ext: str) -> str:
    """Estrae testo dai formati supportati. Mantiene la stessa copertura di
    _read_file_text() di chat_assistant.py. Le immagini ritornano ''."""
    try:
        if ext in {".txt", ".md", ".py", ".json", ".xml", ".csv"}:
            return path.read_text(encoding="utf-8", errors="replace")[:MAX_TEXT_CHARS]
        if ext == ".pdf":
            try:
                import fitz
                doc = fitz.open(str(path))
                return "\n".join(p.get_text() for p in doc)[:MAX_TEXT_CHARS]
            except ImportError:
                import pdfplumber
                with pdfplumber.open(str(path)) as pdf:
                    return "\n".join(p.extract_text() or "" for p in pdf.pages)[:MAX_TEXT_CHARS]
        if ext == ".docx":
            return _docx_full_text(path)[:MAX_TEXT_CHARS]
        if ext in {".xlsx", ".xls"}:
            import openpyxl
            wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
            lines = []
            for ws in wb.worksheets:
                lines.append(f"[Foglio: {ws.title}]")
                for row in ws.iter_rows(values_only=True):
                    lines.append("\t".join(str(c) if c is not None else "" for c in row))
            return "\n".join(lines)[:MAX_TEXT_CHARS]
    except Exception as _e:
        # Nell'app COMPILATA, se un binario nativo (PyMuPDF/lxml/pdfminer) non e'
        # incluso, l'estrazione fallisce qui: logghiamo l'errore reale invece di
        # restituire silenziosamente "" (cosi' l'indice vuoto non resta un mistero).
        try:
            _dbg(f"estrazione FALLITA: {path.name} ({ext}): {type(_e).__name__}: {_e}")
        except Exception:
            pass
        return ""
    return ""


def _iter_files(root: Path):
    """Walk con os.scandir (meno syscall/round-trip SMB). Salta cartelle nascoste
    e di servizio. Robusto: gli errori per-entry non interrompono la scansione."""
    stack = [str(root)]
    while stack:
        d = stack.pop()
        try:
            with os.scandir(d) as it:
                for e in it:
                    try:
                        if e.is_dir(follow_symlinks=False):
                            n = e.name
                            if n.startswith(".") or n in {"__pycache__", "node_modules", ".git"}:
                                continue
                            stack.append(e.path)
                        elif e.is_file(follow_symlinks=False):
                            yield e
                    except OSError:
                        continue
        except OSError:
            continue


def _connect(db_path) -> sqlite3.Connection:
    """Apre una connessione SQLite pronta all'uso. check_same_thread=False +
    busy_timeout: la connessione di LETTURA (thread UI, via _db) e quella di
    SCRITTURA (thread di indicizzazione, in update()) coesistono grazie al WAL,
    senza errori 'same thread' né 'database is locked'."""
    conn = sqlite3.connect(str(db_path), check_same_thread=False)
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA busy_timeout=5000")
    _ensure_schema(conn)
    return conn


def _ensure_schema(conn: sqlite3.Connection):
    conn.execute("CREATE TABLE IF NOT EXISTS meta(path TEXT PRIMARY KEY, sig TEXT)")
    conn.execute("CREATE TABLE IF NOT EXISTS idxinfo(k TEXT PRIMARY KEY, v TEXT)")
    conn.execute(
        "CREATE VIRTUAL TABLE IF NOT EXISTS chunks USING fts5("
        "path UNINDEXED, relpath UNINDEXED, chunk_no UNINDEXED, content, "
        "tokenize='unicode61 remove_diacritics 2')"
    )
    conn.commit()


class FolderIndex:
    def __init__(self, kb_dir: str):
        self.kb_dir = str(Path(kb_dir))
        self.db_path = _db_path_for(self.kb_dir)
        self._conn = None

    # ── connessione di LETTURA (cache, usata dal thread UI per search/count) ──
    def _db(self):
        if self._conn is None:
            self._conn = _connect(self.db_path)
        return self._conn

    def close(self):
        if self._conn is not None:
            try:
                self._conn.close()
            finally:
                self._conn = None

    # ── aggiornamento incrementale (THREAD-SAFE: connessione dedicata) ──
    def update(self, progress=None) -> dict:
        """Indicizza solo i file nuovi/modificati; rimuove i cancellati.
        Apre una connessione DEDICATA, così update() può girare in un QThread
        separato mentre il thread UI interroga l'indice (WAL + busy_timeout).
        `progress(done, total)` opzionale (chiamato ogni ~200 file processati)."""
        t0 = time.time()
        if not fts5_available():
            raise RuntimeError("SQLite FTS5 non disponibile in questo interprete")
        if not Path(self.kb_dir).exists():
            n = 0
            try:
                c = _connect(self.db_path)
                n = c.execute("SELECT COUNT(*) FROM meta").fetchone()[0]
                c.close()
            except Exception:
                pass
            return {"ok": False, "error": "percorso non raggiungibile", "files": n}

        conn = _connect(self.db_path)
        try:
            cur = conn.cursor()
            # Se la LOGICA DI ESTRAZIONE è cambiata (EXTRACTOR_VERSION), l'indice
            # contiene testo vecchio (es. .docx senza tabelle) e i file non
            # risultano modificati → niente re-parse incrementale. Forziamo un
            # rebuild completo una volta sola.
            _ver = cur.execute("SELECT v FROM idxinfo WHERE k='extractor'").fetchone()
            if (_ver[0] if _ver else None) != EXTRACTOR_VERSION:
                _dbg(f"estrattore v{_ver[0] if _ver else '?'} → v{EXTRACTOR_VERSION}: "
                     f"rebuild completo dell'indice (ri-estrazione di tutti i file)")
                cur.execute("DELETE FROM chunks")
                cur.execute("DELETE FROM meta")
                cur.execute("INSERT OR REPLACE INTO idxinfo(k, v) VALUES ('extractor', ?)",
                            (EXTRACTOR_VERSION,))
                conn.commit()
            known = dict(cur.execute("SELECT path, sig FROM meta").fetchall())

            seen, added, changed, skipped = set(), 0, 0, 0
            batch = 0
            root = Path(self.kb_dir)

            for e in _iter_files(root):
                ext = os.path.splitext(e.name)[1].lower()
                if ext not in SUPPORTED_EXT:
                    continue
                try:
                    st = e.stat()
                except OSError:
                    continue
                if st.st_size > MAX_FILE_BYTES:
                    skipped += 1
                    continue
                path = e.path
                seen.add(path)
                sig = f"{int(st.st_mtime)}:{st.st_size}"
                if known.get(path) == sig:
                    continue  # invariato → non riaprire (il guadagno principale)

                text = _extract_text(Path(path), ext)
                try:
                    relpath = str(Path(path).relative_to(root))
                except Exception:
                    relpath = os.path.basename(path)

                cur.execute("DELETE FROM chunks WHERE path = ?", (path,))
                chs = chunk_text(text)
                if chs:
                    cur.executemany(
                        "INSERT INTO chunks(path, relpath, chunk_no, content) VALUES (?,?,?,?)",
                        [(path, relpath, i, c) for i, c in enumerate(chs)],
                    )
                cur.execute("INSERT OR REPLACE INTO meta(path, sig) VALUES (?,?)", (path, sig))
                if path in known:
                    changed += 1
                else:
                    added += 1
                batch += 1
                if batch % 200 == 0:
                    conn.commit()
                    if progress:
                        try: progress(batch, None)
                        except Exception: pass

            # file rimossi dalla cartella → elimina dall'indice
            removed = 0
            for path in list(known.keys()):
                if path not in seen:
                    cur.execute("DELETE FROM chunks WHERE path = ?", (path,))
                    cur.execute("DELETE FROM meta WHERE path = ?", (path,))
                    removed += 1

            conn.commit()
            nfiles = cur.execute("SELECT COUNT(*) FROM meta").fetchone()[0]
            nchunks = cur.execute("SELECT COUNT(*) FROM chunks").fetchone()[0]
            return {"ok": True, "files": nfiles, "chunks": nchunks,
                    "added": added, "changed": changed, "removed": removed,
                    "skipped": skipped, "elapsed": round(time.time() - t0, 2)}
        finally:
            conn.close()

    def count(self):
        try:
            conn = self._db()
            nf = conn.execute("SELECT COUNT(*) FROM meta").fetchone()[0]
            nc = conn.execute("SELECT COUNT(*) FROM chunks").fetchone()[0]
            return nf, nc
        except Exception:
            return 0, 0

    def clear(self):
        try:
            self.close()
            if self.db_path.exists():
                self.db_path.unlink()
            for ext in ("-wal", "-shm"):
                p = Path(str(self.db_path) + ext)
                if p.exists():
                    p.unlink()
        except Exception:
            pass

    # ── ricerca ──
    @staticmethod
    def _build_match(query: str):
        toks = re.findall(r"[0-9A-Za-zÀ-ÿ_]{2,}", (query or "").lower())
        toks = [t for t in toks if t not in _STOP][:12]
        if not toks:
            return None
        # OR di termini fra virgolette: evita che i caratteri speciali rompano
        # la sintassi FTS; i diacritici sono già neutralizzati dal tokenizer.
        return " OR ".join(f'"{t}"' for t in toks)

    def search(self, query: str, top_k: int = 8,
               max_chars_per_chunk: int = 1200, rerank: bool = False):
        """Ricerca BM25 sull'indice. Ritorna (testo_formattato, [(nome, path_assoluto), ...])
        — stesso contratto di vector_db.search(), così si innesta nello stesso canale."""
        match = self._build_match(query)
        if not match:
            return "", []
        try:
            conn = self._db()
            pool = max(top_k, 30) if rerank else top_k
            # snippet(): estrae la finestra di testo ATTORNO ai termini trovati
            # (colonna 3 = content). Senza, un chunk lungo verrebbe mostrato dai
            # primi caratteri e il passaggio pertinente in coda andrebbe perso.
            rows = conn.execute(
                "SELECT path, relpath, content, "
                "snippet(chunks, 3, '', '', ' … ', 64) AS snip, "
                "bm25(chunks) AS rank "
                "FROM chunks WHERE chunks MATCH ? ORDER BY rank LIMIT ?",
                (match, pool),
            ).fetchall()
        except Exception:
            return "", []

        if not rows:
            return "", []

        if rerank and len(rows) > top_k:
            rows = self._semantic_rerank(query, rows, top_k)
        else:
            rows = rows[:top_k]

        parts, sources, seen_src = [], [], set()
        for row in rows:
            path, relpath, content = row[0], row[1], (row[2] or "").strip()
            snip = (row[3] or "").strip() if len(row) > 3 else ""
            if len(content) <= max_chars_per_chunk:
                shown = content
            else:
                # chunk lungo: preferisci la finestra sui match; fallback all'inizio
                shown = snip[:max_chars_per_chunk] if len(snip) > 40 else content[:max_chars_per_chunk]
            parts.append(f"[Fonte: {relpath}]\n{shown}")
            if relpath not in seen_src:
                seen_src.add(relpath)
                sources.append((relpath, path))
        return "\n\n".join(parts), sources

    # ── re-rank semantico opzionale (solo se sentence-transformers c'è) ──
    def _semantic_rerank(self, query, rows, top_k):
        try:
            from sentence_transformers import SentenceTransformer, util
            try:
                from vector_db import resolve_embedding_model
                model_name = resolve_embedding_model("multilingual")
            except Exception:
                model_name = "paraphrase-multilingual-MiniLM-L12-v2"
            global _RERANK_MODEL
            if "_RERANK_MODEL" not in globals() or _RERANK_MODEL is None:
                _RERANK_MODEL = SentenceTransformer(model_name)
            texts = [r[2] or "" for r in rows]
            q_emb = _RERANK_MODEL.encode(query, convert_to_tensor=True, normalize_embeddings=True)
            d_emb = _RERANK_MODEL.encode(texts, convert_to_tensor=True, normalize_embeddings=True)
            scores = util.cos_sim(q_emb, d_emb)[0]
            order = sorted(range(len(rows)), key=lambda i: float(scores[i]), reverse=True)
            return [rows[i] for i in order[:top_k]]
        except Exception:
            return rows[:top_k]


_RERANK_MODEL = None

# ── cache per cartella (come get_vdb) ──
_INDEX_CACHE = {}

def get_index(kb_dir: str) -> "FolderIndex":
    key = str(Path(kb_dir).resolve())
    idx = _INDEX_CACHE.get(key)
    if idx is None:
        idx = FolderIndex(kb_dir)
        _INDEX_CACHE[key] = idx
    return idx


def search_folder(kb_dir: str, query: str, top_k: int = 8, rerank: bool = False):
    """Convenienza: aggiornamento implicito non eseguito qui (lo fa il refresh KB)."""
    return get_index(kb_dir).search(query, top_k=top_k, rerank=rerank)
