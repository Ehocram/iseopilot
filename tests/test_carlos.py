"""
Test incremento "feedback Carlos": (A) recall Conoscenza con inventario di
copertura, (C) template documento personale con bypass del default.
Il renderer chat (B) è client-side: verificato con harness Node in build
(vedi README, sezione Incremento 7); qui si verifica il markup servito.
Esecuzione:
    PYTHONPATH=. APP_DATA_DIR=./data_test python -m pytest tests/ -q
"""
import io
import json
import os
from pathlib import Path

os.environ.setdefault("APP_DATA_DIR", "./data_test")
from cryptography.fernet import Fernet
os.environ.setdefault("APP_SECRET_KEY", Fernet.generate_key().decode())

from fastapi.testclient import TestClient
from app.main import app
from app import store, auth, docgen, knowledge

store.init_db()


def _mk_user(name, dept="IT", is_admin=False):
    store.create_user(name, auth.hash_password("Password123"), dept, is_admin=is_admin)
    c = TestClient(app)
    c.post("/login", data={"username": name, "password": "Password123"},
           follow_redirects=False)
    return c


def _docx_bytes(landscape=False) -> bytes:
    from docx import Document
    from docx.shared import Emu
    d = Document()
    if landscape:
        sec = d.sections[0]
        w, h = sec.page_width, sec.page_height
        sec.page_width, sec.page_height = h, w
    d.add_paragraph("placeholder")
    bio = io.BytesIO()
    d.save(bio)
    return bio.getvalue()


def _pptx_bytes(wide=False) -> bytes:
    from pptx import Presentation
    from pptx.util import Emu
    p = Presentation()
    if wide:
        p.slide_width = Emu(12192000)   # 16:9
        p.slide_height = Emu(6858000)
    bio = io.BytesIO()
    p.save(bio)
    return bio.getvalue()


# ════════════════════════════════════════════════════════════
#  A) Recall Conoscenza: copertura, inventario, diversità
# ════════════════════════════════════════════════════════════
def test_coverage_regex_positivi_e_negativi():
    pos = [
        "tutto quello che abbiamo sulla NIS2",
        "quali certificazioni ISO abbiamo?",
        "abbiamo una policy di backup?",
        "esiste una procedura per gli incidenti?",
        "fammi una panoramica dei documenti privacy",
        "do we have a business continuity plan?",
        "list all documents about GDPR",
        "give me an overview of our security policies",
        "what documents cover supplier audits?",
    ]
    neg = [
        "qual è la scadenza del contratto Rossi?",
        "traduci questo paragrafo in inglese",
        "quanto fattura il fornitore Tecnoplast?",
        "correggi la formula nella cella B2",
    ]
    for q in pos:
        assert knowledge.is_coverage_query(q), q
    for q in neg:
        assert not knowledge.is_coverage_query(q), q


def test_inventory_block_conteggi_dichiarati():
    nomi = [f"doc_{i}.pdf" for i in range(5)]
    out = knowledge._inventory_block("quali documenti abbiamo?", nomi)
    assert "5 documenti" in out and "elenco COMPLETO" in out
    for n in nomi:
        assert n in out
    # troncamento dichiarato, mai silenzioso
    tanti = [f"documento_con_nome_molto_lungo_{i:03d}.docx" for i in range(200)]
    out = knowledge._inventory_block("elenco completo dei documenti", tanti, max_chars=600)
    assert "di 200" in out and "troncato" in out
    # nessun inventario fuori dalle domande di copertura
    assert knowledge._inventory_block("scadenza contratto Rossi", nomi) == ""


class _FakeVdb:
    def __init__(self, nomi, chunks=None, in_doc=True):
        self._nomi = nomi
        self._chunks = chunks or []
        self._in_doc = in_doc

    def list_documents(self):
        return [{"name": n} for n in self._nomi]

    def search_in_document(self, query, nome, n_results=1):
        return f"estratto rilevante da {nome}" if self._in_doc else ""

    def search_raw(self, query, n_results=8):
        return self._chunks[:n_results]

    def search(self, query, n_results=4):
        return "", []


def _patch_vdb(monkeypatch, fake):
    from app.engines import vector_db
    monkeypatch.setattr(knowledge, "kb_available", lambda: True)
    monkeypatch.setattr(vector_db, "get_vdb", lambda cfg: fake)


def test_kb_search_copertura_inventario_diversita(monkeypatch):
    fake = _FakeVdb(
        nomi=["ISO9001.pdf", "ISO14001.pdf", "Policy_Password.docx"],
        chunks=[("c1", "A.docx"), ("c2", "A.docx"), ("c3", "A.docx"), ("c4", "B.docx")],
    )
    _patch_vdb(monkeypatch, fake)
    out, fonti = knowledge.kb_search("IT", "quali certificazioni ISO abbiamo?", n=5)
    # copertura per nome: entrambi i documenti ISO presenti
    assert "[Fonte: ISO9001.pdf]" in out and "[Fonte: ISO14001.pdf]" in out
    # diversità semantica: mai più di 2 passaggi dalla stessa fonte
    assert out.count("[Fonte: A.docx]") <= 2
    # inventario per la domanda di copertura, con conteggio
    assert "[INVENTARIO Conoscenza" in out and "3 documenti" in out
    nomi_fonti = [n for n, _u in fonti]
    assert "ISO9001.pdf" in nomi_fonti and "ISO14001.pdf" in nomi_fonti


def test_kb_search_copertura_senza_chunk_ritorna_inventario(monkeypatch):
    fake = _FakeVdb(nomi=["BCP.docx", "DR_Plan.docx"], chunks=[], in_doc=False)
    _patch_vdb(monkeypatch, fake)
    out, fonti = knowledge.kb_search("IT", "do we have a business continuity plan?", n=5)
    assert "[INVENTARIO Conoscenza" in out and "BCP.docx" in out
    # domanda NON di copertura, niente risultati: comportamento storico (vuoto)
    out2, _ = knowledge.kb_search("IT", "scadenza contratto Tecnoplast", n=5)
    assert out2 == ""


def test_kb_search_dodici_documenti_per_nome_tutti_inclusi(monkeypatch):
    nomi = [f"gdpr_{i:02d}.pdf" for i in range(12)]
    fake = _FakeVdb(nomi=nomi, chunks=[])
    _patch_vdb(monkeypatch, fake)
    out, _ = knowledge.kb_search("IT", "che documenti gdpr abbiamo", n=5)
    inclusi = sum(1 for n in nomi if f"[Fonte: {n}]" in out)
    assert inclusi == 12          # il vecchio tetto era 8
    assert "TUTTI inclusi" in out


# ════════════════════════════════════════════════════════════
#  C) Template personale: validazione, storage, bypass, rotte
# ════════════════════════════════════════════════════════════
def test_validate_template_regole():
    v = docgen.validate_office_template
    assert v(_docx_bytes(), "docx", "mio.docx") == ""
    assert v(_pptx_bytes(), "pptx", "mio.pptx") == ""
    assert "macro" in v(_docx_bytes(), "docx", "mio.docm").lower()
    assert v(b"non uno zip", "docx", "x.docx") != ""
    # zip con macro dentro: rifiutato anche con estensione lecita
    import zipfile
    bio = io.BytesIO()
    with zipfile.ZipFile(bio, "w") as z:
        z.writestr("[Content_Types].xml", "<Types></Types>")
        z.writestr("word/document.xml", "<w:document/>")
        z.writestr("word/vbaProject.bin", "MACRO")
    assert "macro" in v(bio.getvalue(), "docx", "furbo.docx").lower()
    bio = io.BytesIO()
    with zipfile.ZipFile(bio, "w") as z:
        z.writestr("[Content_Types].xml",
                   "<Types><Override ContentType='application/vnd.ms-word.document.macroEnabled.main+xml'/></Types>")
        z.writestr("word/document.xml", "<w:document/>")
    assert "macro" in v(bio.getvalue(), "docx", "furbo2.docx").lower()
    # docx spacciato per pptx: part attesa assente
    assert "ppt/presentation.xml" in v(_docx_bytes(), "pptx", "x.pptx")


def test_template_storage_roundtrip_e_isolamento():
    assert docgen.save_user_template("tplA@test", "docx", _docx_bytes(), "Solution Centre.docx") == ""
    t = docgen.get_user_template("tplA@test", "docx")
    assert t and t[1] == "Solution Centre.docx" and Path(t[0]).is_file()
    assert docgen.get_user_template("tplB@test", "docx") is None   # isolamento
    st = docgen.user_templates_status("tplA@test")
    assert st["docx"]["name"] == "Solution Centre.docx" and st["pptx"] is None
    assert docgen.delete_user_template("tplA@test", "docx") is True
    assert docgen.get_user_template("tplA@test", "docx") is None


def test_gen_docx_bypass_template_personale(tmp_path):
    # prova di bypass verificabile: template PERSONALE in orizzontale ->
    # il documento generato eredita l'orientamento del template personale
    custom = tmp_path / "landscape.docx"
    custom.write_bytes(_docx_bytes(landscape=True))
    spec = {"title": "Prova", "sections": [{"heading": "S1", "paragraphs": ["p"],
                                            "bullets": ["b1"]}]}
    from docx import Document
    path_custom, _ = docgen.gen_docx(spec, template_path=str(custom))
    path_default, _ = docgen.gen_docx(spec)
    w_custom = Document(path_custom).sections[0].page_width
    h_custom = Document(path_custom).sections[0].page_height
    w_default = Document(path_default).sections[0].page_width
    assert w_custom > h_custom              # orizzontale: viene dal template personale
    assert w_custom != w_default            # e differisce dal default ISEO


def test_gen_docx_template_corrotto_fail_loud(tmp_path):
    rotto = tmp_path / "rotto.docx"
    rotto.write_text("non sono un docx")
    import pytest
    with pytest.raises(ValueError) as e:
        docgen.gen_docx({"title": "x", "sections": []}, template_path=str(rotto))
    assert "non è utilizzabile" in str(e.value)


def test_gen_pptx_bypass_e_fail_loud(tmp_path):
    wide = tmp_path / "wide.pptx"
    wide.write_bytes(_pptx_bytes(wide=True))
    from pptx import Presentation
    spec = {"title": "Deck", "slides": [{"title": "S1", "bullets": ["a", "b"]}]}
    path, _ = docgen.gen_pptx(spec, template_path=str(wide))
    assert Presentation(path).slide_width == 12192000   # eredita il 16:9 del personale
    rotto = tmp_path / "rotto.pptx"
    rotto.write_text("no")
    import pytest
    with pytest.raises(ValueError):
        docgen.gen_pptx(spec, template_path=str(rotto))


def test_generate_instrada_template_per_formato(monkeypatch, tmp_path):
    visto = {}
    monkeypatch.setattr(docgen, "build_spec",
                        lambda fmt, req, ctx, st, hist="", note_utente="": {"title": "T", "sections": [], "slides": []})
    monkeypatch.setattr(docgen, "gen_docx",
                        lambda spec, template_path=None: visto.setdefault("docx", template_path) or (str(tmp_path / "o.docx"), "o.docx"))
    monkeypatch.setattr(docgen, "_docx_to_pdf", lambda p: str(tmp_path / "o.pdf"))
    (tmp_path / "o.pdf").write_bytes(b"%PDF")
    docgen.generate("pdf", "req", "", {}, templates={"docx": "/percorso/mio.docx"})
    assert visto["docx"] == "/percorso/mio.docx"   # il PDF usa il template Word personale


def test_api_template_upload_status_delete_audit():
    c = _mk_user("tplapi@test")
    # stato iniziale vuoto
    st = c.get("/api/template").json()
    assert st == {"docx": None, "pptx": None}
    # upload valido
    r = c.post("/api/template",
               files={"file": ("Solution Centre.docx", _docx_bytes(),
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    j = r.json()
    assert j["ok"] is True and j["status"]["docx"]["name"] == "Solution Centre.docx"
    # macro rifiutata con motivo
    r = c.post("/api/template", files={"file": ("evil.docm", _docx_bytes(), "application/octet-stream")})
    j = r.json()
    assert j["ok"] is False and "macro" in j["error"].lower()
    # estensione non ammessa
    r = c.post("/api/template", files={"file": ("x.txt", b"ciao", "text/plain")})
    assert r.json()["ok"] is False
    # rimozione
    r = c.post("/api/template/delete", data={"fmt": "docx"})
    assert r.json()["ok"] is True
    assert c.get("/api/template").json()["docx"] is None
    # audit tracciato
    azioni = {r["action"] for r in store.audit_query(username="tplapi@test")}
    assert {"template_caricato", "template_rifiutato", "template_rimosso"} <= azioni


def test_api_template_richiede_login():
    c = TestClient(app)
    assert c.get("/api/template").status_code == 401
    assert c.post("/api/template/delete", data={"fmt": "docx"}).status_code == 401


def test_generazione_in_chat_usa_template_personale(monkeypatch, tmp_path):
    c = _mk_user("tplchat@test")
    r = c.post("/api/template", files={"file": ("SC.docx", _docx_bytes(), "application/octet-stream")})
    assert r.json()["ok"] is True
    atteso = docgen.get_user_template("tplchat@test", "docx")[0]

    visto = {}
    out = tmp_path / "gen.docx"
    out.write_bytes(_docx_bytes())

    def fake_generate(fmt, req, ctx, st, hist="", templates=None, note_utente=""):
        visto["fmt"], visto["templates"] = fmt, dict(templates or {})
        return str(out), "gen.docx"

    from app import main as main_mod
    monkeypatch.setattr(main_mod.docgen, "generate", fake_generate)
    r = c.post("/api/chat", json={"messages": [{"role": "user", "content":
               "creami un documento word sulla sicurezza"}],
               "engine": "claude", "free_mode": True})
    assert r.status_code == 200
    assert visto["fmt"] == "docx"
    assert visto["templates"]["docx"] == atteso        # bypass instradato
    assert "sul tuo template **SC.docx**" in r.text    # dichiarato all'utente
    # pulizia: rimuovi il template per non influenzare altri test
    c.post("/api/template/delete", data={"fmt": "docx"})


def _dotx_bytes() -> bytes:
    """Modello Word (.dotx): stesso package del .docx con content-type
    'template' — quello che python-docx rifiuta e che va normalizzato."""
    import zipfile
    raw = _docx_bytes()
    bio_in = io.BytesIO(raw)
    bio_out = io.BytesIO()
    with zipfile.ZipFile(bio_in) as zi, zipfile.ZipFile(bio_out, "w") as zo:
        for item in zi.infolist():
            data = zi.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.decode("utf-8").replace(
                    "wordprocessingml.document.main+xml",
                    "wordprocessingml.template.main+xml").encode("utf-8")
            zo.writestr(item, data)
    return bio_out.getvalue()


def test_normalize_template_dotx_e_idempotenza():
    dotx = _dotx_bytes()
    from docx import Document
    import pytest
    with pytest.raises(Exception):
        Document(io.BytesIO(dotx))                    # il modello viene rifiutato
    conv = docgen.normalize_office_template(dotx)
    assert conv != dotx
    Document(io.BytesIO(conv))                        # il normalizzato si apre
    assert docgen.normalize_office_template(conv) == conv   # idempotente sui .docx


def test_upload_dotx_accettato_e_normalizzato():
    c = _mk_user("tpldotx@test")
    r = c.post("/api/template",
               files={"file": ("Solution_Center.dotx", _dotx_bytes(), "application/octet-stream")})
    j = r.json()
    assert j["ok"] is True and j["status"]["docx"]["name"] == "Solution_Center.dotx"
    # su disco è già un documento apribile: la generazione non fallirà
    path, _nome = docgen.get_user_template("tpldotx@test", "docx")
    from docx import Document
    Document(path)
    c.post("/api/template/delete", data={"fmt": "docx"})


def test_gen_docx_personale_usa_stili_del_template(tmp_path):
    """Contratto sui template personali: Heading 1 per le sezioni, Subtitle per
    il sottotitolo, e MAI bullet persi (numbering reale oppure glifo •)."""
    import zipfile
    tpl = tmp_path / "personale.docx"
    tpl.write_bytes(_docx_bytes())
    spec = {"title": "T", "subtitle": "S",
            "sections": [{"heading": "Sez", "paragraphs": ["p"], "bullets": ["b1", "b2"]}]}
    path, _ = docgen.gen_docx(spec, template_path=str(tpl))
    with zipfile.ZipFile(path) as z:
        doc = z.read("word/document.xml").decode("utf-8")
    assert 'w:pStyle w:val="Heading1"' in doc
    assert 'w:pStyle w:val="Subtitle"' in doc
    assert ("<w:numPr>" in doc) or ("• b1" in doc)     # elenco vero o glifo, mai piatto


# ── Attività (Cowork): documento garantito e template personale ──
def _cowork_events(monkeypatch, uid, domanda, fake_complete, tpl_path=None):
    from app import cowork
    monkeypatch.setattr(cowork.knowledge, "kb_list",
                        lambda dept: [{"name": "NIS2 Guide.docx"}])
    monkeypatch.setattr(cowork.knowledge, "kb_search",
                        lambda dept, q, n=4: ("frammento sulla NIS2", []))
    monkeypatch.setattr(cowork, "complete", fake_complete)
    return list(cowork.run("IT", uid, domanda, {"claude_api_key": "k"}, []))


def test_cowork_finalizza_documento_a_passi_esauriti(monkeypatch):
    """Caso Carlos: il planner spende tutti i passi in ricerca e non arriva mai
    a 'componi' — il documento va COMPOSTO comunque in finalizzazione."""
    import json as _json

    def fake_complete(system, user, settings, max_tokens=2500, timeout=120):
        if "JSON valido" in system:      # chiamata di finalizzazione: la spec
            return _json.dumps({"title": "NIS2 Sales Guide",
                                "sections": [{"heading": "S", "paragraphs": ["p"],
                                              "bullets": ["b"]}]})
        if "LIMITE PASSI RAGGIUNTO" in user:   # sintesi finale testuale
            return "Guida pronta, file al download."
        return '{"azione": "cerca", "query": "nis2"}'   # loop: sempre ricerca

    events = _cowork_events(monkeypatch, "cw_fin@test",
                            "genera un documento word: guida NIS2 per i venditori",
                            fake_complete)
    blob = "".join(events)
    assert '"kind": "download"' in blob                 # il file c'è
    assert "documento composto in finalizzazione" in blob
    assert "sintesi parziale dichiarata" not in blob


def test_cowork_senza_richiesta_documento_niente_finalizzazione(monkeypatch):
    def fake_complete(system, user, settings, max_tokens=2500, timeout=120):
        if "LIMITE PASSI RAGGIUNTO" in user:
            return "Sintesi testuale."
        return '{"azione": "cerca", "query": "nis2"}'

    events = _cowork_events(monkeypatch, "cw_nofin@test",
                            "riassumi cosa dice la NIS2 sui fornitori",
                            fake_complete)
    blob = "".join(events)
    assert '"kind": "download"' not in blob             # nessun file inventato
    assert "sintesi parziale dichiarata" in blob        # chiusura onesta invariata


def test_cowork_componi_usa_template_personale(monkeypatch):
    import json as _json
    assert docgen.save_user_template("cw_tpl@test", "docx", _docx_bytes(), "SC.docx") == ""
    atteso = docgen.get_user_template("cw_tpl@test", "docx")[0]
    visto = {}
    orig = docgen.gen_docx

    def spia(spec, template_path=None):
        visto["tpl"] = template_path
        return orig(spec, template_path=template_path)

    from app import cowork
    monkeypatch.setattr(cowork.docgen, "gen_docx", spia)
    calls = {"n": 0}

    def fake_complete(system, user, settings, max_tokens=2500, timeout=120):
        calls["n"] += 1
        if calls["n"] == 1:
            return _json.dumps({"azione": "componi",
                                "spec": {"title": "T", "sections":
                                         [{"heading": "S", "paragraphs": ["p"], "bullets": []}]}})
        return '{"azione": "rispondi", "testo": "Fatto, file pronto. [Fonte: NIS2 Guide.docx]"}'

    events = _cowork_events(monkeypatch, "cw_tpl@test",
                            "crea un word di prova", fake_complete)
    blob = "".join(events)
    assert visto["tpl"] == atteso                       # bypass anche in Attività
    assert "template: SC.docx" in blob                  # dichiarato nel footer
    docgen.delete_user_template("cw_tpl@test", "docx")


# ── Lingua: la richiesta esplicita PREVALE sull'impostazione ──
def test_build_system_precedenza_lingua_esplicita():
    from app.orchestrator import build_system
    out = build_system("Aziendale formale", "Italiano")
    assert "SOLO in italiano" in out          # il default resta
    assert "PREVALE" in out                   # ...ma la richiesta esplicita vince
    out_auto = build_system("Aziendale formale", "Auto")
    assert "PREVALE" not in out_auto          # con Auto non serve l'eccezione


def test_build_spec_dichiara_lingua_e_prevalenza(monkeypatch):
    catturato = {}

    def fake_complete(system, user, settings, max_tokens=4000, timeout=120):
        catturato["system"] = system
        return '{"title": "x", "sections": []}'

    monkeypatch.setattr(docgen.orchestrator, "complete", fake_complete)
    docgen.build_spec("docx", "redacta una especificación técnica", "",
                      {"reply_lang": "Spagnolo", "claude_api_key": "k"})
    assert "Spagnolo" in catturato["system"]
    assert "PREVALE" in catturato["system"]


def test_cowork_riceve_conversazione_e_regola_lingua(monkeypatch):
    catture = []

    def fake_complete(system, user, settings, max_tokens=2500, timeout=120):
        catture.append((system, user))
        return '{"azione": "rispondi", "testo": "Hecho. [Fonte: NIS2 Guide.docx]"}'

    from app import cowork
    monkeypatch.setattr(cowork.knowledge, "kb_list",
                        lambda dept: [{"name": "NIS2 Guide.docx"}])
    monkeypatch.setattr(cowork, "complete", fake_complete)
    events = list(cowork.run("IT", "cw_lang@test", "no, la volevo in spagnolo",
                             {"claude_api_key": "k", "reply_lang": "Italiano"}, [],
                             conversazione="UTENTE: bozza di specifica tecnica del cilindro\nASSISTENTE: Ecco la bozza…"))
    system, user = catture[0]
    assert "LINGUA" in system and "prevale" in system.lower()
    assert "[CONVERSAZIONE PRECEDENTE" in user
    assert "bozza di specifica tecnica del cilindro" in user   # la correzione ha il suo contesto
    assert any('"type": "done"' in e for e in events)


def test_api_chat_cowork_passa_gli_ultimi_turni(monkeypatch):
    store.set_setting("cowork_enabled", "1")
    try:
        c = _mk_user("cw_hist@test")
        catturato = {}

        def fake_run(dept, uid, domanda, settings, anon_names, conversazione="", note_utente=""):
            catturato["conv"] = conversazione
            catturato["domanda"] = domanda
            yield 'data: {"type": "done"}\n\n'

        from app import cowork
        monkeypatch.setattr(cowork, "run", fake_run)
        r = c.post("/api/chat", json={"mode": "cowork", "engine": "claude",
                   "messages": [
                       {"role": "user", "content": "prepara la specifica tecnica X100"},
                       {"role": "assistant", "content": "Ecco la specifica in italiano…"},
                       {"role": "user", "content": "no, la volevo in spagnolo"}]})
        assert r.status_code == 200
        assert catturato["domanda"] == "no, la volevo in spagnolo"
        assert "prepara la specifica tecnica X100" in catturato["conv"]
        assert "UTENTE:" in catturato["conv"] and "ASSISTENTE:" in catturato["conv"]
        assert "in spagnolo" not in catturato["conv"]   # l'ultimo messaggio è il COMPITO
    finally:
        store.set_setting("cowork_enabled", "0")


# ── B) markup servito per il renderer/typography ────────────
def test_chat_page_serve_hook_renderer_e_template():
    c = _mk_user("uihook@test")
    html = c.get("/").text
    assert 'id="tpl-btn"' in html and 'id="tpl-bar"' in html
    css = c.get("/static/app.css").text
    assert ".code-copy" in css and ".bubble h2" in css.replace(".msg .bubble h1,.msg .bubble h2", ".bubble h2")
    js = c.get("/static/chat.js").text
    assert "renderMarkdownish" in js and "escapeHtml(blk.code)" in js


if __name__ == "__main__":
    import inspect
    import sys
    import tempfile

    class _MP:
        def setattr(self, obj, name, value):
            setattr(obj, name, value)

    fns = [f for n, f in sorted(globals().items())
           if n.startswith("test_") and inspect.isfunction(f)]
    failed = 0
    for f in fns:
        try:
            kwargs = {}
            sig = inspect.signature(f)
            td = None
            if "tmp_path" in sig.parameters:
                td = tempfile.TemporaryDirectory()
                kwargs["tmp_path"] = Path(td.name)
            if "monkeypatch" in sig.parameters:
                kwargs["monkeypatch"] = _MP()
            f(**kwargs)
            print(f"  PASS  {f.__name__}")
            if td:
                td.cleanup()
        except Exception as e:
            failed += 1
            print(f"  FAIL  {f.__name__}: {e}")
    print(f"\n{len(fns)-failed}/{len(fns)} test superati.")
    sys.exit(1 if failed else 0)
